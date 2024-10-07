import streamlit as st
from fpdf import FPDF
from docx import Document
from PIL import Image
import io
from tools import load_yolo_model
from PIL import Image
import supervision as sv
import numpy as np



favicon = Image.open("/workspaces/canprev-modules/CanPrev_4D-logo.png")
st.set_page_config(
    page_title="Canprev AI",  # Title of the web page
    page_icon=favicon,                # Favicon, use emoji or a file path
    layout="wide",                       # Layout options: "wide" or "centered"
    initial_sidebar_state="expanded"     # Initial state of the sidebar: "auto", "expanded", or "collapsed"
)
st.image('CanPrev_4D-logo.png', width=300)
st.markdown("""
<style>

	.stTabs [data-baseweb="tab-list"] {
		gap: 3px;
    }

	.stTabs [data-baseweb="tab"] {
		height: 40px;
        white-space: pre-wrap;
		background-color: #13276F;
		border-radius: 4px 4px 0px 0px;
		gap: 1px;
		padding: 10px 2px 10px 2px;
        color: white;
    }

	.stTabs [aria-selected="true"] {
  		background-color: #FFFFFF;
        color: #13276F;
        border: 2px solid #13276F;
        border-bottom: none;
	}

</style>""", unsafe_allow_html=True)

# Load Model
model_side_QA = load_yolo_model('models/model_side_view_qa.pt')
# model_bottle_seal = load_yolo_model('models/bottleseal_nano_model.pt')
# model_unopened_side_view_checklist = load_yolo_model()
# model_liquid_powder = load_yolo_model('models/model_powder_liquid_lump.pt')
# model_bottle_dent = load_yolo_model('models/model_bottle_dent-50.pt')
# model_bottle_label = load_yolo_model('models/model_bottle_label-50.pt')
# model_bottle_cap_checklist = load_yolo_model('models/cap_condition_checklist.pt')


from collections import defaultdict

CHECKLIST = defaultdict(list)
THINGS_TO_CHECK = {'label', 'botle_with_neckband', 'curved_shoulder'}

def update_CHECKLIST(key, value):
    CHECKLIST[key].append(value)


def detect_top_view(image):
    return True



def detect_side_view(image, view_name, model):
    image_array = Image.open(image)

    result = model(image_array)[0]
    detections = sv.Detections.from_ultralytics(result)
    detections[detections.confidence > .60]
    for thing in THINGS_TO_CHECK:
        if thing in detections.data['class_name']:
            update_CHECKLIST(thing, True)

    return True

def detect_bottom_view(image):
    return {
        "OCR Text": "LOT123, Exp: 2025-12-31, Price: $20, Material Type: Plastic"
    }



st.title("Product Inspection Dashboard")
st.subheader("Upload Images")

# First row
col1, col2 = st.columns(2)
with col1:
    top_view_img = st.file_uploader("Top View", type=["jpg", "png", "jpeg"])
with col2:
    bottom_view_img = st.file_uploader("Bottom View", type=["jpg", "png", "jpeg"])

# Second row
col3, col4 = st.columns(2)
with col3:
    left_view_img = st.file_uploader("Left View", type=["jpg", "png", "jpeg"])
with col4:
    right_view_img = st.file_uploader("Right View", type=["jpg", "png", "jpeg"])

# Third row
col5, col6 = st.columns(2)
with col5:
    front_view_img = st.file_uploader("Front View", type=["jpg", "png", "jpeg"])
with col6:
    back_view_img = st.file_uploader("Back View", type=["jpg", "png", "jpeg"])


st.divider()


def display_checklist(results, view_name):
    cols = st.columns(4)
    idx = 0
    for key, value in results.items():
        col = cols[idx % 4]
        value = all(value)
        if value:
            col.checkbox(f"{key}: {value}", value=True, key=f"{view_name}_{key}")
        else:
            col.checkbox(f"{key}: {value}", value=False, key=f"{view_name}_{key}")
        idx += 1


def merge_side_view_analysis(images):
    
    for view_name, image in images.items():
        if image:
            view_results = detect_side_view(image, view_name, model=model_side_QA)
    return True

all_results = {}
if top_view_img:
    st.subheader("Unopened Bottle Checklist")
    top_view_results = detect_top_view(top_view_img)
    all_results["Top View"] = top_view_results
    display_checklist(top_view_results, "Top")

if bottom_view_img:
    st.subheader("Bottom View Analysis")
    bottom_view_results = detect_bottom_view(bottom_view_img)
    all_results["Bottom View"] = bottom_view_results
    display_checklist(bottom_view_results, "Bottom")


side_images = {
    "Left": left_view_img,
    "Right": right_view_img,
    "Front": front_view_img,
    "Back": back_view_img
}

# side_view_results = merge_side_view_analysis(side_images)
if left_view_img and right_view_img and front_view_img and back_view_img:
    st.subheader("Side View Analysis (Left, Right, Front, Back)")
    display_checklist(CHECKLIST, "Side")

def generate_pdf(results):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", size=12)
    
    for view_name, items in results.items():
        pdf.cell(200, 10, text=f"{view_name} Analysis", ln=True, align='C')
        pdf.ln(10)

        # Table Header
        pdf.set_font("Arial", style='B', size=12)
        pdf.cell(100, 10, text="Feature", border=1)
        pdf.cell(100, 10, text="Value", border=1)
        pdf.ln()

        # Table Content
        pdf.set_font("Arial", size=12)
        for key, value in items.items():
            pdf.cell(100, 10, text=key, border=1)
            pdf.cell(100, 10, text=str(value), border=1)
            pdf.ln()
        pdf.ln(10)
    
    return pdf


def generate_docx(results):
    doc = Document()
    doc.add_heading('Product Inspection Report', 0)

    for view_name, items in results.items():
        doc.add_heading(f"{view_name} Analysis", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Value'

        for key, value in items.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

    return doc


download_enabled = all([top_view_img, bottom_view_img, any(side_images.values())])

pdf_download_button, docs_download_button = st.columns([1,1], vertical_alignment='bottom')
with pdf_download_button:
    if download_enabled:
        pdf = generate_pdf(all_results)
        pdf_output = io.BytesIO()
        pdf.output(pdf_output)
        pdf_output.seek(0)

        st.download_button(label="Download PDF", data=pdf_output, file_name="inspection_report.pdf", mime="application/pdf")

with docs_download_button:
    if download_enabled:
        doc = generate_docx(all_results)
        doc_output = io.BytesIO()
        doc.save(doc_output)
        doc_output.seek(0)

        st.download_button(label="Download DOCX", data=doc_output, file_name="inspection_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

